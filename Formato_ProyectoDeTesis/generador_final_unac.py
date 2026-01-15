import json
import os
import platform
import subprocess
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

class SistemasHenyerEngine:
    def __init__(self, json_path):
        with open(json_path, 'r', encoding='utf-8') as f:
            self.data = json.load(f)
        self.doc = Document()
        self.conf = self.data['configuracion']

    def set_cell_background(self, cell, color_hex):
        """Pinta el fondo de una celda"""
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color_hex)
        tcPr.append(shd)

    def insertar_tabla_encabezado(self):
        """Inserta la tabla de control IDIE de la UNAC"""
        table = self.doc.add_table(rows=4, cols=5)
        table.style = 'Table Grid'
        table.autofit = False
        
        for row in table.rows:
            row.cells[0].width = Cm(3.0)
            for i in range(1, 5): row.cells[i].width = Cm(3.4)

        celda_logo = table.cell(0, 0).merge(table.cell(3, 0))
        p_logo = celda_logo.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if os.path.exists(self.conf['ruta_logo']):
            p_logo.add_run().add_picture(self.conf['ruta_logo'], width=Cm(2.0))
        else:
            p_logo.add_run("LOGO\nUNAC").bold = True

        celda_idie = table.cell(0, 1).merge(table.cell(0, 4))
        celda_idie.text = "I + D + i + e"
        celda_idie.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        celda_idie.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        celda_idie.paragraphs[0].runs[0].bold = True
        self.set_cell_background(celda_idie, self.conf['color_encabezado'])

        table.cell(1, 1).text = "PROCESO NIVEL 0:"
        table.cell(1, 2).merge(table.cell(1, 4)).text = "INVESTIGACIÓN CIENTÍFICA Y TECNOLÓGICA"
        table.cell(2, 1).text = "REGISTRO"
        table.cell(2, 2).merge(table.cell(2, 4)).text = "PROYECTO DE INVESTIGACIÓN - TESIS"
        
        table.cell(3, 1).text = "Código: M.IDIE.01/R4"
        table.cell(3, 2).text = "Versión: 01"
        table.cell(3, 3).text = "Fecha: 14/12/2023"
        table.cell(3, 4).text = "Página: 1 de 10"

        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = self.conf['fuente_tabla']
                        run.font.size = Pt(self.conf['tamano_tabla'])

        self.doc.add_paragraph().paragraph_format.space_after = Pt(10)

    def aplicar_estilos_base(self):
        style = self.doc.styles['Normal']
        style.font.name = self.conf['fuente_normal']
        style.font.size = Pt(self.conf['tamano_normal'])

    def construir(self):
        self.aplicar_estilos_base()
        for i, pag in enumerate(self.data['paginas']):
            if i > 0: self.doc.add_page_break()
            self.insertar_tabla_encabezado()

            if pag['tipo'] == 'caratula' or pag['tipo'] == 'lista':
                p = self.doc.add_paragraph(pag['titulo'])
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(14)
                for item in pag['items']:
                    self.doc.add_paragraph(item)

            elif pag['tipo'] == 'indice':
                p = self.doc.add_paragraph(pag['titulo'])
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].bold = True
                for item in pag['items']:
                    para = self.doc.add_paragraph(item['texto'])
                    para.paragraph_format.left_indent = Cm(item['indent'] * 0.7)
                    if item['bold']: para.runs[0].bold = True

            elif pag['tipo'] == 'contenido_detallado':
                for idx_cap, cap in enumerate(pag['capitulos']):
                    if idx_cap > 0:
                        self.doc.add_page_break()
                        self.insertar_tabla_encabezado()
                    t = self.doc.add_paragraph(cap['titulo'])
                    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = t.runs[0]
                    run.bold = True
                    run.font.size = Pt(14)
                    for sec in cap['secciones']:
                        if sec['sub']:
                            st = self.doc.add_paragraph(sec['sub'])
                            st.runs[0].bold = True
                        tp = self.doc.add_paragraph(sec['texto'])
                        tp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        output = self.conf['nombre_archivo']
        self.doc.save(output)
        print(f"\n✓ Documento '{output}' generado exitosamente.")
        self.abrir_archivo(output)

    def abrir_archivo(self, ruta):
        try:
            if platform.system() == 'Windows': os.startfile(ruta)
            elif platform.system() == 'Darwin': subprocess.run(['open', ruta])
            else: subprocess.run(['xdg-open', ruta])
        except: pass

if __name__ == "__main__":
    folder = 'Formato_ProyectoDeTesis'
    
    print("========================================")
    print("      SISTEMAS-HENYER DOC GENERATOR     ")
    print("========================================")
    print("Seleccione el tipo de proyecto a generar:")
    print("1. Enfoque Cuantitativo (proyecto_tesis_cuantitativo.json)")
    print("2. Enfoque Cualitativo (proyecto_tesis_cualitativo.json)")
    print("========================================")
    
    opcion = input("Ingrese 1 o 2: ")
    
    if opcion == "1":
        file = 'proyecto_tesis_cuantitativo.json'
    elif opcion == "2":
        file = 'proyecto_tesis_cualitativo.json'
    else:
        print("Opción no válida. Cerrando programa.")
        exit()

    ruta_completa = os.path.join(folder, file)

    if not os.path.exists(ruta_completa):
        print(f"ERROR: No se encontró el archivo en: {ruta_completa}")
    else:
        print(f"Procesando: {file}...")
        engine = SistemasHenyerEngine(ruta_completa)
        engine.construir()