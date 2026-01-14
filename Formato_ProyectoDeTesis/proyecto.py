from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import platform
import subprocess

# --- 1. FUNCIONES DE DISEÑO Y COLOR ---

def set_cell_background(cell, color_hex):
    """Pinta el fondo de una celda (Hex sin #)"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def configurar_estilos(doc):
    """Times New Roman 12 para todo"""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Títulos en Negro y Negrita
    for i in range(1, 4):
        estilo = doc.styles[f'Heading {i}']
        estilo.font.name = 'Times New Roman'
        estilo.font.size = Pt(12)
        estilo.font.color.rgb = RGBColor(0, 0, 0)
        estilo.font.bold = True
        estilo.paragraph_format.space_before = Pt(12)
        estilo.paragraph_format.space_after = Pt(12)

# --- 2. FUNCIÓN GENERADORA DEL CUADRO (ENCABEZADO SIMULADO) ---

def insertar_tabla_encabezado(doc):
    """Inserta la tabla tipo encabezado EN EL CUERPO del documento"""
    # Crear tabla de 4 filas y 5 columnas
    table = doc.add_table(rows=4, cols=5)
    table.style = 'Table Grid'
    table.autofit = False
    
    # Ancho total aprox de la hoja A4 menos márgenes
    ancho_total = Inches(6.5) 
    
    # Configurar anchos: Col 0 (Logo) más ancha
    for row in table.rows:
        row.cells[0].width = Cm(3.0)
        for i in range(1, 5):
            row.cells[i].width = Cm(3.4)

    # --- FUSIONES Y CONTENIDO ---
    
    # A. LOGO (Columna 0, Filas 0-3)
    celda_logo = table.cell(0, 0).merge(table.cell(3, 0))
    p_logo = celda_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    ruta_logo = "Formato_ProyectoDeTesis/logo_unac.png" # <--- TU IMAGEN AQUI
    if os.path.exists(ruta_logo):
        try:
            run_logo = p_logo.add_run()
            run_logo.add_picture(ruta_logo, width=Cm(2.2))
        except:
            p_logo.add_run("LOGO\nUNAC").bold = True
    else:
        p_logo.add_run("LOGO\nUNAC").bold = True

    # B. BARRA AZUL SUPERIOR (Fila 0, Cols 1-4)
    celda_idie = table.cell(0, 1).merge(table.cell(0, 4))
    celda_idie.text = "I + D + i + e"
    celda_idie.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    celda_idie.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    celda_idie.paragraphs[0].runs[0].bold = True
    set_cell_background(celda_idie, "BDD7EE") # Color Azul Claro

    # C. PROCESO (Fila 1)
    celda_p1 = table.cell(1, 1)
    celda_p1.text = "PROCESO NIVEL 0:"
    celda_p2 = table.cell(1, 2).merge(table.cell(1, 4))
    celda_p2.text = "INVESTIGACIÓN CIENTÍFICA Y TECNOLÓGICA"

    # D. REGISTRO (Fila 2)
    celda_r1 = table.cell(2, 1)
    celda_r1.text = "REGISTRO"
    celda_r2 = table.cell(2, 2).merge(table.cell(2, 4))
    celda_r2.text = "PROYECTO DE INVESTIGACIÓN - TESIS"

    # E. DATOS INFERIORES (Fila 3)
    table.cell(3, 1).text = "Código: M.IDIE.01/R4"
    table.cell(3, 2).text = "Versión: 01"
    table.cell(3, 3).text = "Fecha: 14/12/2023"
    table.cell(3, 4).text = "Página: 1  de  10"

    # F. FORMATO DE TEXTO DE LA TABLA (Arial 9pt)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                # Alinear izquierda excepto el título azul
                if cell == celda_idie:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif cell == celda_logo:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(9)

    # Espacio después de la tabla para que no se pegue al título
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

def nueva_pagina_con_encabezado(doc):
    """Crea salto de página e inserta inmediatamente el cuadro"""
    doc.add_page_break()
    insertar_tabla_encabezado(doc)

# --- 3. FUNCIONES DE TEXTO ---

def agregar_titulo(doc, texto):
    p = doc.add_paragraph(texto)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.bold = True
    run.font.size = Pt(14)
    p.paragraph_format.space_after = Pt(12)

def agregar_subtitulo(doc, texto):
    p = doc.add_paragraph(texto)
    run = p.runs[0]
    run.bold = True
    p.paragraph_format.space_before = Pt(6)

def agregar_texto(doc, texto):
    p = doc.add_paragraph(texto)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# --- 4. BLOQUE PRINCIPAL ---

def crear_tesis_completa():
    doc = Document()
    configurar_estilos(doc)

    # ================= PÁGINA 1: CARÁTULA =================
    insertar_tabla_encabezado(doc) # Insertar al inicio sin salto de página

    for _ in range(2): doc.add_paragraph()
    
    p = doc.add_paragraph('CARÁTULA')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)
    
    items_caratula = [
        "● TÍTULO (Contiene las variables a investigar)",
        "● AUTOR(ES)",
        "● LUGAR Y FECHA",
        "● PÁGINA DE RESPETO"
    ]
    for item in items_caratula:
        p = doc.add_paragraph(item)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)

    # ================= PÁGINA 2: INFORMACIÓN BÁSICA =================
    nueva_pagina_con_encabezado(doc)
    
    agregar_titulo(doc, "INFORMACIÓN BÁSICA")
    items_info = [
        "FACULTAD", "UNIDAD DE INVESTIGACIÓN", "AUTOR (es) / CÓDIGO ORCID / DNI",
        "ASESOR y CO ASESOR", "LUGAR DE EJECUCIÓN", "UNIDAD DE ANÁLISIS",
        "TIPO / ENFOQUE / DISEÑO", "TEMA OCDE"
    ]
    for item in items_info:
        doc.add_paragraph(f"● {item}")

    # ================= PÁGINA 3: ÍNDICE COMPLETO =================
    nueva_pagina_con_encabezado(doc)
    
    agregar_titulo(doc, "ÍNDICE")
    p_desc = doc.add_paragraph("(contenido, tablas, figuras, abreviaturas)")
    p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # LISTA DE ÍNDICE CORREGIDA Y COMPLETA
    items_indice = [
        "INTRODUCCIÓN",
        "I. PLANTEAMIENTO DEL PROBLEMA",
        "   1.1. Realidad problemática",
        "   1.2. Formulación del problema",
        "   1.3. Objetivos",
        "   1.4. Justificación",
        "   1.5. Delimitantes",
        "II. MARCO TEÓRICO",
        "   2.1. Antecedentes",
        "   2.2. Bases teóricas",
        "   2.3. Marco conceptual",
        "   2.4. Definición de términos",
        "III. HIPÓTESIS Y VARIABLES",
        "   3.1. Hipótesis",
        "   3.1.1. Operacionalización",
        "IV. METODOLOGÍA DEL PROYECTO",
        "   4.1. Diseño metodológico",
        "   4.2. Método de investigación",
        "   4.3. Población y muestra",
        "   4.5. Técnicas e instrumentos",
        "   4.7. Aspectos Éticos",
        "V. CRONOGRAMA",
        "VI. PRESUPUESTO",
        "VII. REFERENCIAS",
        "VIII. ANEXOS"
    ]
    
    for item in items_indice:
        p = doc.add_paragraph(item)
        # Tabuladores simples para simular índice
        if item.startswith("   "):
            p.paragraph_format.left_indent = Cm(1)
        else:
            p.runs[0].bold = True

    # ================= PÁGINA 4: INTRODUCCIÓN =================
    nueva_pagina_con_encabezado(doc)
    agregar_titulo(doc, "INTRODUCCIÓN")
    agregar_texto(doc, "Breve presentación del contenido y relevancia del proyecto.")

    # ================= PÁGINA 5: CAPÍTULO I =================
    nueva_pagina_con_encabezado(doc)
    agregar_titulo(doc, "I. PLANTEAMIENTO DEL PROBLEMA")
    agregar_texto(doc, "Capítulo dedicado a la fundamentación de la problemática.")

    agregar_subtitulo(doc, "1.1. Realidad problemática")
    agregar_texto(doc, "Seleccionar el problema en el contexto de las líneas de investigación prioritarias o transversales.")
    
    agregar_subtitulo(doc, "1.2. Formulación del problema")
    agregar_texto(doc, "Definir el problema general y específicos.")

    agregar_subtitulo(doc, "1.3. Objetivos")
    agregar_texto(doc, "Establecer el objetivo general y los objetivos específicos.")

    agregar_subtitulo(doc, "1.4. Justificación")
    agregar_texto(doc, "Sustentar de acuerdo a la naturaleza del problema.")

    agregar_subtitulo(doc, "1.5. Delimitantes")
    agregar_texto(doc, "Definir los límites de la investigación: teórica, temporal y espacial.")

    # ================= PÁGINA 6: CAPÍTULO II =================
    nueva_pagina_con_encabezado(doc)
    agregar_titulo(doc, "II. MARCO TEÓRICO")
    
    agregar_subtitulo(doc, "2.1. Antecedentes")
    agregar_texto(doc, "Reportes de investigaciones previas a nivel Internacional y nacional.")

    agregar_subtitulo(doc, "2.2. Bases teóricas")
    agregar_texto(doc, "Exponer y argumentar las teorías sustantivas del problema a investigar.")

    agregar_subtitulo(doc, "2.3. Marco conceptual")
    agregar_texto(doc, "Elaborar nuevos constructos fundamentados de las teorías.")

    agregar_subtitulo(doc, "2.4. Definición de términos")
    agregar_texto(doc, "Términos funcionales a la investigación del problema.")

    # ================= PÁGINA 7: CAPÍTULO III =================
    nueva_pagina_con_encabezado(doc)
    agregar_titulo(doc, "III. HIPÓTESIS Y VARIABLES")
    
    agregar_subtitulo(doc, "3.1. Hipótesis")
    agregar_texto(doc, "Planteamiento de la hipótesis general y las específicas.")

    agregar_subtitulo(doc, "3.1.1. Operacionalización")
    agregar_texto(doc, "Definición conceptual y operacional de variables, dimensiones, indicadores, items.")

    # ================= PÁGINA 8: CAPÍTULO IV =================
    nueva_pagina_con_encabezado(doc)
    agregar_titulo(doc, "IV. METODOLOGÍA DEL PROYECTO")
    
    agregar_subtitulo(doc, "4.1. Diseño metodológico")
    agregar_texto(doc, "Se concreta según la naturaleza del problema a investigar.")

    agregar_subtitulo(doc, "4.2. Método de investigación")
    agregar_texto(doc, "Básica, Aplicada, Filosófica o Humanística.")

    agregar_subtitulo(doc, "4.3. Población y muestra")
    agregar_texto(doc, "Definición del universo y el grupo de estudio.")

    agregar_subtitulo(doc, "4.5. Técnicas e instrumentos")
    agregar_texto(doc, "Herramientas para la recolección de la información.")

    agregar_subtitulo(doc, "4.7. Aspectos Éticos")
    agregar_texto(doc, "Consideraciones éticas aplicadas en la investigación.")

    # ================= PÁGINA 9: CAPÍTULOS RESTANTES (V - VIII) =================
    nueva_pagina_con_encabezado(doc)
    
    agregar_titulo(doc, "V. CRONOGRAMA")
    agregar_texto(doc, "Calendarización en meses mediante Diagrama de Gantt.")
    doc.add_paragraph()

    agregar_titulo(doc, "VI. PRESUPUESTO")
    agregar_texto(doc, "Financiamiento y costo real en soles o moneda equivalente.")
    doc.add_paragraph()

    agregar_titulo(doc, "VII. REFERENCIAS")
    agregar_texto(doc, "Uso de normas (APA, Vancouver, ISO 690 o IEEE).")
    doc.add_paragraph()

    agregar_titulo(doc, "VIII. ANEXOS")
    agregar_texto(doc, "Matriz de consistencia y propuesta de instrumentos.")

    # --- GUARDAR ---
    nombre_archivo = 'Proyecto_UNAC_Final.docx'
    doc.save(nombre_archivo)
    print(f'✓ Documento generado con Encabezado Sólido e Índice Completo: {nombre_archivo}')
    
    # ABRIR
    try:
        if platform.system() == 'Windows': os.startfile(nombre_archivo)
        elif platform.system() == 'Darwin': subprocess.run(['open', nombre_archivo])
        else: subprocess.run(['xdg-open', nombre_archivo])
    except: pass

if __name__ == '__main__':
    crear_tesis_completa()