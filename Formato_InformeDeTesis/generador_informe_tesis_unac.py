import json
import os
import platform
import subprocess
import sys
from docx import Document
from docx.shared import Cm, Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
FORMATS_DIR = os.path.join(BASE_DIR, "formats")
IMAGENES_DIR = os.path.join(BASE_DIR, "Imagenes")

def cargar_contenido(path_archivo):
    if not os.path.exists(path_archivo):
        nombre = os.path.basename(path_archivo)
        path_archivo = os.path.join(FORMATS_DIR, nombre)
        
    if not os.path.exists(path_archivo):
        raise FileNotFoundError(f"No se encontro el JSON: {path_archivo}")

    with open(path_archivo, 'r', encoding='utf-8') as f:
        return json.load(f)

def configurar_formato_unac(doc):
    for section in doc.sections:
        section.page_width = Cm(21.0); section.page_height = Cm(29.7)
        section.left_margin = Cm(3.5); section.right_margin = Cm(2.5)
        section.top_margin = Cm(3.0); section.bottom_margin = Cm(3.0)

    style = doc.styles['Normal']
    style.font.name = 'Arial'; style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def agregar_bloque(doc, texto, negrita=False, tamano=12, antes=0, despues=0, cursiva=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE 
    p.paragraph_format.space_before = Pt(antes)
    p.paragraph_format.space_after = Pt(despues)
    run = p.add_run(texto)
    run.bold = negrita; run.italic = cursiva; run.font.size = Pt(tamano)
    return p

def agregar_titulo_formal(doc, texto, espaciado_antes=0):
    h = doc.add_heading(level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run(texto)
    run.font.name = 'Arial'; run.font.size = Pt(14); run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    h.paragraph_format.space_before = Pt(espaciado_antes)
    h.paragraph_format.space_after = Pt(12)

def agregar_nota_guia(doc, texto):
    if not texto: return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(f"Nota: {texto}")
    run.font.name = 'Arial'; run.font.size = Pt(10); run.italic = True
    run.font.color.rgb = RGBColor(89, 89, 89) 
    p.paragraph_format.space_after = Pt(12)

def crear_caratula_dinamica(doc, data):
    c = data['caratula']
    agregar_bloque(doc, c['universidad'], negrita=True, tamano=18, despues=4)
    agregar_bloque(doc, c['facultad'], negrita=True, tamano=14, despues=4)
    agregar_bloque(doc, c['escuela'], negrita=True, tamano=14, despues=25)

    ruta_logo = os.path.join(IMAGENES_DIR, "LogoUNAC.png")
    if os.path.exists(ruta_logo):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.add_run().add_picture(ruta_logo, width=Inches(3.2))
    else:
        agregar_bloque(doc, "[LOGO INSTITUCIONAL]", tamano=10, antes=40, despues=40)

    agregar_bloque(doc, c['tipo_documento'], negrita=True, tamano=16, antes=30)
    agregar_bloque(doc, c['titulo_placeholder'], negrita=True, tamano=14, antes=30, despues=30)
    agregar_bloque(doc, c['frase_grado'], tamano=12, antes=10)
    agregar_bloque(doc, c['grado_objetivo'], negrita=True, tamano=13, despues=35)
    
    agregar_bloque(doc, c['label_autor'], negrita=True, tamano=12, antes=5)
    agregar_bloque(doc, c['label_asesor'], negrita=True, tamano=12, antes=5, despues=20)
    agregar_bloque(doc, c['label_linea'], tamano=11, cursiva=True, despues=40)
    agregar_bloque(doc, c['fecha'], tamano=12); agregar_bloque(doc, c['pais'], negrita=True, tamano=12)

def agregar_preliminares_dinamico(doc, data):
    p = data['preliminares']
    doc.add_paragraph(); doc.add_page_break() 
    if 'dedicatoria' in p:
        agregar_titulo_formal(doc, p['dedicatoria']['titulo'])
        doc.add_paragraph(p['dedicatoria']['texto']).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_page_break()
    if 'resumen' in p:
        agregar_titulo_formal(doc, p['resumen']['titulo'])
        doc.add_paragraph(p['resumen']['texto']).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_page_break()
    idx = p['indices']
    agregar_titulo_formal(doc, idx['contenido'])
    doc.add_paragraph("(Generar Indice Automatico)").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    if 'introduccion' in p:
        agregar_titulo_formal(doc, p['introduccion']['titulo'])
        doc.add_paragraph(p['introduccion']['texto']).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_page_break()

def agregar_cuerpo_dinamico(doc, data):
    for cap in data['cuerpo']:
        h = doc.add_heading(level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = h.add_run(cap['titulo'])
        run.font.name = 'Arial'; run.font.size = Pt(14); run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        h.paragraph_format.space_before = Pt(24); h.paragraph_format.space_after = Pt(18)
        if 'contenido' in cap:
            for item in cap['contenido']:
                sub = doc.add_paragraph()
                run_sub = sub.add_run(item['texto'])
                run_sub.font.name = 'Arial'; run_sub.font.size = Pt(12); run_sub.bold = True
        doc.add_page_break()

def agregar_finales_dinamico(doc, data):
    fin = data['finales']
    agregar_titulo_formal(doc, fin['referencias']['titulo'])
    doc.add_page_break()
    agregar_titulo_formal(doc, fin['anexos']['titulo_seccion'])
    doc.add_page_break()

def agregar_numeracion_paginas(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; p.clear() 
        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)
        run.font.name = 'Arial'; run.font.size = Pt(10)

def generar_documento_core(ruta_json, ruta_salida):
    data = cargar_contenido(ruta_json)
    doc = Document()
    configurar_formato_unac(doc)
    crear_caratula_dinamica(doc, data)
    agregar_preliminares_dinamico(doc, data)
    agregar_cuerpo_dinamico(doc, data)
    agregar_finales_dinamico(doc, data)
    agregar_numeracion_paginas(doc)
    doc.save(ruta_salida)
    # CORRECCION: Emoji quitado
    print(f"[OK] Generado: {ruta_salida}")
    return ruta_salida

if __name__ == "__main__":
    if len(sys.argv) > 2:
        path_json_arg = sys.argv[1]
        path_output_arg = sys.argv[2]
        try:
            generar_documento_core(path_json_arg, path_output_arg)
        except Exception as e:
            # CORRECCION: Emoji quitado
            print(f"[ERROR] Error en generador: {str(e)}")
            sys.exit(1)
    else:
        print("="*40)
        print("   GENERADOR DE INFORME (CLI)")
        print("="*40)
        print("1. Enfoque CUANTITATIVO")
        print("2. Enfoque CUALITATIVO")
        
        try: opcion = input(">> Opcion: ").strip()
        except: sys.exit()

        if opcion == '1':
            json_file = "unac_informe_cuant.json"
            out_file = "Informe_Cuantitativo.docx"
        elif opcion == '2':
            json_file = "unac_informe_cual.json"
            out_file = "Informe_Cualitativo.docx"
        else:
            print("Opcion invalida."); sys.exit()
            
        json_path = os.path.join(FORMATS_DIR, json_file)
        
        try:
            ruta = generar_documento_core(json_path, out_file)
            if platform.system() == 'Windows': os.startfile(ruta)
        except Exception as e:
            print(f"Error: {e}")