import tkinter as tk
from tkinter import messagebox, ttk
import json
import os
import platform
import subprocess
import threading
from docx import Document
from docx.shared import Cm, Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ==========================================
# 1. LÓGICA DE GENERACIÓN (BACKEND)
# ==========================================

def cargar_contenido(nombre_archivo_json):
    ruta_script = os.path.dirname(__file__)
    # CORRECCIÓN: Aseguramos que busque en la carpeta 'formats' subiendo un nivel
    ruta_json = os.path.join(ruta_script, "..", "formats", nombre_archivo_json)
    ruta_absoluta = os.path.abspath(ruta_json)
    
    if not os.path.exists(ruta_absoluta):
        raise FileNotFoundError(f"No se encontró el archivo: {ruta_absoluta}")

    with open(ruta_absoluta, 'r', encoding='utf-8') as f:
        return json.load(f)

def configurar_formato_unac(doc):
    """Configura márgenes y fuente base Arial 12"""
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(3.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(3.0)
        section.bottom_margin = Cm(3.0)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    rFonts = style.element.rPr.rFonts
    rFonts.set(qn('w:ascii'), 'Arial')
    rFonts.set(qn('w:hAnsi'), 'Arial')
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def agregar_bloque(doc, texto, negrita=False, tamano=12, antes=0, despues=0, cursiva=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE 
    p.paragraph_format.space_before = Pt(antes)
    p.paragraph_format.space_after = Pt(despues)
    run = p.add_run(texto)
    run.bold = negrita
    run.italic = cursiva
    run.font.size = Pt(tamano)
    return p

def agregar_titulo_formal(doc, texto, espaciado_antes=0):
    h = doc.add_heading(level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run(texto)
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    h.paragraph_format.space_before = Pt(espaciado_antes)
    h.paragraph_format.space_after = Pt(12)

def agregar_nota_guia(doc, texto):
    if not texto: return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(f"Nota: {texto}")
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(89, 89, 89) 
    p.paragraph_format.space_after = Pt(12)

def crear_caratula_dinamica(doc, data):
    c = data['caratula']
    agregar_bloque(doc, c['universidad'], negrita=True, tamano=18, despues=4)
    agregar_bloque(doc, c['facultad'], negrita=True, tamano=14, despues=4)
    agregar_bloque(doc, c['escuela'], negrita=True, tamano=14, despues=25)

    ruta_script = os.path.dirname(__file__)
    ruta_logo = os.path.join(ruta_script, "..", "Imagenes", "LogoUNAC.png")
    
    if os.path.exists(ruta_logo):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_logo = p_logo.add_run()
        run_logo.add_picture(ruta_logo, width=Inches(3.2))
    else:
        agregar_bloque(doc, "[LOGO INSTITUCIONAL]", tamano=10, antes=40, despues=40)

    agregar_bloque(doc, c['tipo_documento'], negrita=True, tamano=16, antes=30)
    agregar_bloque(doc, c['titulo_placeholder'], negrita=True, tamano=14, antes=30, despues=30)
    agregar_bloque(doc, c['frase_grado'], tamano=12, antes=10)
    agregar_bloque(doc, c['grado_objetivo'], negrita=True, tamano=13, despues=35)
    
    agregar_bloque(doc, c['label_autor'], negrita=True, tamano=12, antes=5)
    agregar_bloque(doc, c['label_asesor'], negrita=True, tamano=12, antes=5, despues=20)
    agregar_bloque(doc, c['label_linea'], tamano=11, cursiva=True, despues=40)
    
    agregar_bloque(doc, c['fecha'], tamano=12)
    agregar_bloque(doc, c['pais'], negrita=True, tamano=12)

def agregar_preliminares_dinamico(doc, data):
    p = data['preliminares']
    
    doc.add_paragraph() 
    doc.add_page_break() 

    if 'dedicatoria' in p:
        agregar_titulo_formal(doc, p['dedicatoria']['titulo'])
        doc.add_paragraph(p['dedicatoria']['texto']).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_page_break()

    if 'resumen' in p:
        agregar_titulo_formal(doc, p['resumen']['titulo'])
        if 'nota' in p['resumen']:
            p_nota = doc.add_paragraph()
            p_nota.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_nota = p_nota.add_run(p['resumen']['nota'])
            run_nota.italic = True
            run_nota.font.size = Pt(11)
        doc.add_paragraph(p['resumen']['texto']).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_page_break()

    idx = p['indices']
    agregar_titulo_formal(doc, idx['contenido'])
    doc.add_paragraph(idx.get('placeholder', '(Generarlo)')).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if 'tablas' in idx:
        agregar_titulo_formal(doc, idx['tablas'], espaciado_antes=30)
        doc.add_paragraph(idx.get('placeholder', '(Generarlo)')).alignment = WD_ALIGN_PARAGRAPH.CENTER

    if 'figuras' in idx:
        agregar_titulo_formal(doc, idx['figuras'], espaciado_antes=30)
        doc.add_paragraph(idx.get('placeholder', '(Generarlo)')).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if 'abreviaturas' in idx:
        agregar_titulo_formal(doc, idx['abreviaturas'], espaciado_antes=30)
        doc.add_paragraph(idx.get('placeholder', '(Generarlo)')).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    if 'introduccion' in p:
        agregar_titulo_formal(doc, p['introduccion']['titulo'])
        par_intro = doc.add_paragraph()
        par_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        par_intro.add_run(p['introduccion']['texto'])
        doc.add_page_break()

def agregar_cuerpo_dinamico(doc, data):
    caps = data['cuerpo']

    for cap in caps:
        h = doc.add_heading(level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = h.add_run(cap['titulo'])
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        h.paragraph_format.space_before = Pt(24)
        h.paragraph_format.space_after = Pt(18)

        if 'nota_capitulo' in cap:
             agregar_nota_guia(doc, cap['nota_capitulo'])

        if 'contenido' in cap:
            for item in cap['contenido']:
                sub = doc.add_paragraph()
                run_sub = sub.add_run(item['texto'])
                run_sub.font.name = 'Arial'
                run_sub.font.size = Pt(12)
                run_sub.bold = True
                sub.paragraph_format.space_before = Pt(12)
                sub.paragraph_format.space_after = Pt(6)
                
                if 'nota' in item:
                    agregar_nota_guia(doc, item['nota'])
        
        doc.add_page_break()

def agregar_finales_dinamico(doc, data):
    fin = data['finales']

    ref = fin['referencias']
    agregar_titulo_formal(doc, ref['titulo'])
    if 'nota' in ref: agregar_nota_guia(doc, ref['nota'])
    if 'ejemplo' in ref: doc.add_paragraph(ref['ejemplo'])
    doc.add_page_break()

    anx = fin['anexos']
    agregar_titulo_formal(doc, anx['titulo_seccion'])

    for anexo in anx['lista']:
        p_anexo = doc.add_paragraph()
        run_anx = p_anexo.add_run(anexo['titulo'])
        run_anx.bold = True
        
        if 'nota' in anexo:
            agregar_nota_guia(doc, anexo['nota'])
        
        if 'tabla_headers' in anexo:
            table = doc.add_table(rows=1, cols=len(anexo['tabla_headers']))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for idx, txt in enumerate(anexo['tabla_headers']):
                hdr_cells[idx].text = txt
                hdr_cells[idx].paragraphs[0].runs[0].bold = True
            
        doc.add_paragraph()
    
    doc.add_page_break()

def agregar_numeracion_paginas(doc):
    """
    Agrega numeración de páginas en la parte INFERIOR DERECHA.
    Según el PDF normativo: 'Numeración de páginas: inferior derecho'.
    """
    for section in doc.sections:
        footer = section.footer
        # Limpiamos el párrafo existente si hay alguno para evitar conflictos
        if footer.paragraphs:
            p = footer.paragraphs[0]
            p.clear()
        else:
            p = footer.add_paragraph()
            
        # CORRECCIÓN CLAVE: Alineación a la DERECHA
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        run = p.add_run()
        # Campo complejo para el número de página
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        
        # Formato de la fuente del número
        run.font.name = 'Arial'
        run.font.size = Pt(10) # Estándar limpio

# --- Lógica Principal del Hilo ---
def logica_generacion(tipo, label_status, root):
    try:
        label_status.config(text="⏳ Generando estructura...", fg="#f1c40f")
        
        if tipo == "CUANTI":
            nombre_json = "tesis_content_cuantitativo.json"
            nombre_salida = "Tesis_Cuantitativa_UNAC.docx"
        else:
            nombre_json = "tesis_content_cualitativa.json"
            nombre_salida = "Tesis_Cualitativa_UNAC.docx"

        data = cargar_contenido(nombre_json)
        doc = Document()
        
        configurar_formato_unac(doc)
        crear_caratula_dinamica(doc, data)
        agregar_preliminares_dinamico(doc, data)
        agregar_cuerpo_dinamico(doc, data)
        agregar_finales_dinamico(doc, data)
        
        # IMPORTANTE: La numeración se agrega al final para aplicarse a todas las secciones
        agregar_numeracion_paginas(doc)
        
        ruta_final = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", nombre_salida))
        
        try:
            doc.save(ruta_final)
        except PermissionError:
            messagebox.showerror("Error", "⚠️ El archivo Word está abierto.\nPor favor, ciérralo e intenta de nuevo.")
            label_status.config(text="❌ Archivo abierto bloqueado", fg="red")
            return

        label_status.config(text="✅ ¡Generado con Éxito!", fg="#2ecc71")
        
        if platform.system() == 'Windows':
            os.startfile(ruta_final)
        else:
            cmd = 'open' if platform.system() == 'Darwin' else 'xdg-open'
            subprocess.call((cmd, ruta_final))
        
        # Cierra la ventana después de 1 segundo para una experiencia fluida
        root.after(1000, root.destroy) 
            
    except Exception as e:
        messagebox.showerror("Error Crítico", f"Ocurrió un error:\n{e}")
        label_status.config(text="❌ Error crítico", fg="red")

# ==========================================
# 2. INTERFAZ GRÁFICA (FRONTEND ELEGANTE)
# ==========================================

def iniciar_app():
    root = tk.Tk()
    root.title("Generador de Tesis UNAC")
    root.geometry("500x650")
    root.resizable(False, False)
    
    COLOR_FONDO = "#1a253a"
    COLOR_TEXTO = "#ecf0f1"
    COLOR_DORADO = "#f39c12"
    COLOR_BTN_BG = "#34495e"
    
    root.configure(bg=COLOR_FONDO)

    main_frame = tk.Frame(root, bg=COLOR_FONDO)
    main_frame.pack(expand=True, fill="both", padx=20, pady=20)

    lbl_titulo = tk.Label(main_frame, text="INFORME DE TESIS\nPREGRADO", 
                          font=("Helvetica", 22, "bold"), 
                          bg=COLOR_FONDO, fg=COLOR_TEXTO, justify="center")
    lbl_titulo.pack(pady=(20, 5))

    lbl_sub = tk.Label(main_frame, text="UNAC", 
                       font=("Times New Roman", 20, "bold"), 
                       bg=COLOR_FONDO, fg=COLOR_DORADO, justify="center")
    lbl_sub.pack(pady=(0, 20))

    ruta_script = os.path.dirname(__file__)
    ruta_logo_png = os.path.join(ruta_script, "..", "Imagenes", "LogoUNAC.png")
    
    logo_frame = tk.Frame(main_frame, bg=COLOR_FONDO, height=150)
    logo_frame.pack(pady=10)

    try:
        img_raw = tk.PhotoImage(file=ruta_logo_png)
        img = img_raw.subsample(2, 2)
        lbl_img = tk.Label(logo_frame, image=img, bg=COLOR_FONDO)
        lbl_img.image = img 
        lbl_img.pack()
    except:
        lbl_ph = tk.Label(logo_frame, text="[ LOGO UNAC ]", 
                          font=("Arial", 12, "bold"), fg=COLOR_DORADO, bg=COLOR_FONDO,
                          relief="groove", borderwidth=2, width=20, height=5)
        lbl_ph.pack()

    tk.Frame(main_frame, bg=COLOR_DORADO, height=2, width=300).pack(pady=20)

    def on_enter(e): e.widget['background'] = COLOR_DORADO; e.widget['foreground'] = "#000"
    def on_leave(e): e.widget['background'] = COLOR_BTN_BG; e.widget['foreground'] = "#fff"

    def crear_boton_elegante(texto, funcion):
        btn = tk.Button(main_frame, text=texto, font=("Segoe UI", 11, "bold"),
                        bg=COLOR_BTN_BG, fg="white", 
                        activebackground=COLOR_DORADO, activeforeground="black",
                        relief="flat", cursor="hand2", width=30, height=2,
                        command=lambda: threading.Thread(target=funcion).start())
        btn.bind("<Enter>", on_enter)
        btn.bind("<Leave>", on_leave)
        btn.pack(pady=10)
        return btn

    crear_boton_elegante("ENFOQUE CUANTITATIVO", lambda: logica_generacion("CUANTI", lbl_status, root))
    crear_boton_elegante("ENFOQUE CUALITATIVO", lambda: logica_generacion("CUALI", lbl_status, root))

    lbl_status = tk.Label(main_frame, text="Sistema listo.", 
                          font=("Arial", 10), bg=COLOR_FONDO, fg="#95a5a6")
    lbl_status.pack(side="bottom", pady=20)

    root.mainloop()

if __name__ == "__main__":
    iniciar_app()