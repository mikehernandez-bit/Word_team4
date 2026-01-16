import json
import os
import sys
import platform
import subprocess
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

class SistemasHenyerEngine:
    def __init__(self, json_path):
        self.base_dir = os.path.dirname(os.path.abspath(__file__))
        
        # Validación de ruta del JSON
        if not os.path.isabs(json_path):
            json_path = os.path.join(self.base_dir, json_path)

        if not os.path.exists(json_path):
            raise FileNotFoundError(f"CRITICO: No se encontro el archivo de configuracion: {json_path}")
            
        with open(json_path, 'r', encoding='utf-8') as f:
            self.data = json.load(f)
        
        self.doc = Document()
        # Configuración por defecto
        self.conf = self.data.get('configuracion', {
            "fuente_normal": "Arial", 
            "tamano_normal": 11,
            "ruta_logo": "assets/LogoUNAC.png", # Nombre por defecto
            "color_encabezado": "D9D9D9",
            "fuente_tabla": "Arial Narrow",
            "tamano_tabla": 9
        })

    def _resolve_asset_path(self, filename_from_json):
        """
        Busca la imagen de forma inteligente:
        1. Intenta la ruta tal cual viene del JSON.
        2. Si falla, limpia la ruta y busca solo el nombre del archivo en la carpeta actual.
        3. Si falla, intenta un nombre por defecto 'LogoUNAC.png'.
        """
        # Opción A: Ruta combinada (base_dir + ruta_json)
        # Esto fallaba antes porque duplicaba la carpeta 'Formato_ProyectoDeTesis'
        path_A = os.path.join(self.base_dir, filename_from_json)
        if os.path.exists(path_A):
            return path_A
            
        # Opción B: Limpieza de ruta (SOLUCIÓN CLAVE)
        # Si el JSON dice 'Carpeta/foto.png', extraemos solo 'foto.png' y buscamos aquí.
        clean_name = os.path.basename(filename_from_json)
        path_B = os.path.join(self.base_dir, clean_name)
        if os.path.exists(path_B):
            return path_B

        # Opción C: Fallback total
        path_C = os.path.join(self.base_dir, "assets", "LogoUNAC.png")
        if os.path.exists(path_C):
            return path_C
            
        return path_A # Retornamos el original para que falle y sepamos por qué

    def set_cell_background(self, cell, color_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        clean_hex = color_hex.replace("#", "")
        shd.set(qn('w:fill'), clean_hex)
        tcPr.append(shd)

    def add_field(self, paragraph, instr_text):
        run = paragraph.add_run()
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = instr_text
        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_char_begin)
        run._r.append(instr)
        run._r.append(fld_char_end)

    def insertar_tabla_encabezado(self):
        table = self.doc.add_table(rows=4, cols=5)
        table.style = 'Table Grid'
        table.autofit = False
        
        for row in table.rows:
            row.cells[0].width = Cm(3.0)
            for i in range(1, 5): row.cells[i].width = Cm(3.4)

        # --- LOGO ---
        celda_logo = table.cell(0, 0).merge(table.cell(3, 0))
        p_logo = celda_logo.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 1. Obtenemos nombre del JSON o usamos default
        nombre_logo = self.conf.get('ruta_logo', 'assets/LogoUNAC.png')
        # 2. Resolvemos la ruta absoluta con la nueva lógica inteligente
        ruta_logo = self._resolve_asset_path(nombre_logo)

        if os.path.exists(ruta_logo):
            run_img = p_logo.add_run()
            # Ajustamos el ancho para que quepa bien en la celda
            run_img.add_picture(ruta_logo, width=Cm(2.2))
        else:
            # Debug visual en el Word si falla
            p_logo.add_run("LOGO\nNO ENCONTRADO").bold = True

        # --- TITULO I+D+i+e ---
        celda_idie = table.cell(0, 1).merge(table.cell(0, 4))
        celda_idie.text = "I + D + i + e"
        celda_idie.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        celda_idie.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        celda_idie.paragraphs[0].runs[0].bold = True
        self.set_cell_background(celda_idie, self.conf.get('color_encabezado', 'D9D9D9'))

        # --- RESTO DE LA TABLA ---
        table.cell(1, 1).text = "PROCESO NIVEL 0:"
        table.cell(1, 2).merge(table.cell(1, 4)).text = "INVESTIGACION CIENTIFICA Y TECNOLOGICA"
        table.cell(2, 1).text = "REGISTRO"
        table.cell(2, 2).merge(table.cell(2, 4)).text = "PROYECTO DE INVESTIGACION - TESIS"
        table.cell(3, 1).text = "Codigo: M.IDIE.01/R4"
        table.cell(3, 2).text = "Version: 01"
        table.cell(3, 3).text = "Fecha: 14/12/2023"
        page_cell = table.cell(3, 4)
        page_para = page_cell.paragraphs[0]
        for run in page_para.runs:
            run.text = ""
        page_para.add_run("Pagina: ")
        self.add_field(page_para, "PAGE")
        page_para.add_run(" de ")
        self.add_field(page_para, "NUMPAGES")

        fuente_tbl = self.conf.get('fuente_tabla', 'Arial Narrow')
        tamano_tbl = self.conf.get('tamano_tabla', 9)
        
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = fuente_tbl
                        run.font.size = Pt(tamano_tbl)
        
        self.doc.add_paragraph().paragraph_format.space_after = Pt(10)

    def aplicar_estilos_base(self):
        style = self.doc.styles['Normal']
        style.font.name = self.conf.get('fuente_normal', 'Arial')
        style.font.size = Pt(self.conf.get('tamano_normal', 11))
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    def construir(self, output_path):
        self.aplicar_estilos_base()
        paginas = self.data.get('paginas', [])
        
        for i, pag in enumerate(paginas):
            if i > 0: self.doc.add_page_break()
            self.insertar_tabla_encabezado()
            
            tipo = pag.get('tipo', 'generico')
            
            if tipo in ['caratula', 'lista']:
                if pag.get('titulo'):
                    p = self.doc.add_paragraph(pag['titulo'])
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.runs[0].bold = True; p.runs[0].font.size = Pt(14)
                for item in pag.get('items', []): self.doc.add_paragraph(item)

            elif tipo == 'indice':
                p = self.doc.add_paragraph(pag.get('titulo', 'INDICE'))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].bold = True
                for item in pag.get('items', []):
                    para = self.doc.add_paragraph(item.get('texto', ''))
                    para.paragraph_format.left_indent = Cm(item.get('indent', 0) * 0.7)
                    if item.get('bold'): para.runs[0].bold = True

            elif tipo == 'contenido_detallado':
                for idx_cap, cap in enumerate(pag.get('capitulos', [])):
                    if idx_cap > 0: self.doc.add_paragraph()
                    if cap.get('titulo'):
                        t = self.doc.add_paragraph(cap['titulo'])
                        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        t.runs[0].bold = True; t.runs[0].font.size = Pt(12)
                    for sec in cap.get('secciones', []):
                        if sec.get('sub'): 
                            st = self.doc.add_paragraph(sec['sub'])
                            st.runs[0].bold = True
                        if sec.get('texto'):
                            self.doc.add_paragraph(sec['texto']).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        full_output_path = os.path.abspath(output_path)
        self.doc.save(full_output_path)
        print(f"[OK] Documento generado: {full_output_path}")

if __name__ == "__main__":
    if len(sys.argv) > 2:
        json_path_arg = sys.argv[1]
        output_path_arg = sys.argv[2]
        
        try:
            engine = SistemasHenyerEngine(json_path_arg)
            engine.construir(output_path_arg)
        except Exception as e:
            print(f"[ERROR] Error en generador: {str(e)}")
            sys.exit(1)

    else:
        BASE_DIR = os.path.dirname(os.path.abspath(__file__))
        FORMATS_DIR = os.path.join(BASE_DIR, "formats")
        
        print("="*40)
        print("   GENERADOR DE PROYECTO (CLI)")
        print("="*40)
        print("1. Enfoque CUANTITATIVO")
        print("2. Enfoque CUALITATIVO")
        
        try: opcion = input(">> Ingrese opcion (1/2): ").strip()
        except: sys.exit()
        
        if opcion == "1":
            json_name = "unac_proyecto_cuant.json"
            out_name = "Proyecto_Cuantitativo.docx"
        elif opcion == "2":
            json_name = "unac_proyecto_cual.json"
            out_name = "Proyecto_Cualitativo.docx"
        else:
            print("Opcion no valida."); sys.exit()
            
        json_path = os.path.join(FORMATS_DIR, json_name)
        
        if os.path.exists(json_path):
            try:
                engine = SistemasHenyerEngine(json_path)
                engine.construir(out_name)
                if platform.system() == 'Windows': os.startfile(out_name)
            except Exception as e:
                print(f"Error: {e}")
        else:
            print(f"No se encontro el JSON: {json_path}")
