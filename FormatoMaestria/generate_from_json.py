import os
import json
import platform
import subprocess
import sys
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# -------------------------
# UTILIDADES JSON / PATHS
# -------------------------

def load_json(path: str) -> dict:
    if not os.path.exists(path):
        raise FileNotFoundError(f"JSON no encontrado: {path}")
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)

def resolve_path(base_dir: str, maybe_rel: str) -> str:
    return maybe_rel if os.path.isabs(maybe_rel) else os.path.join(base_dir, maybe_rel)

def resolve_config_path(base_dir: str, config_path: str) -> str:
    if os.path.isabs(config_path):
        return config_path
    candidate = os.path.join(base_dir, config_path)
    if os.path.exists(candidate):
        return candidate
    return os.path.abspath(config_path)

# -------------------------
# WORD: SETUP / HELPERS
# -------------------------

def open_document(path: str):
    try:
        if platform.system() == "Windows":
            os.startfile(path)
        elif platform.system() == "Darwin":
            subprocess.run(["open", path], check=False)
        else:
            subprocess.run(["xdg-open", path], check=False)
    except Exception as exc:
        print(f"[WARN] No se pudo abrir el documento: {exc}")

def set_page_setup(doc: Document, cfg: dict):
    page_setup = cfg.get("page_setup", {})
    margins = page_setup.get("margins_cm", {})

    sec = doc.sections[0]
    sec.page_width = Mm(210); sec.page_height = Mm(297)

    sec.left_margin = Cm(float(margins.get("left", 3.5)))
    sec.right_margin = Cm(float(margins.get("right", 2.5)))
    sec.top_margin = Cm(float(margins.get("top", 3.0)))
    sec.bottom_margin = Cm(float(margins.get("bottom", 3.0)))

    font_cfg = page_setup.get("font", {})
    font_name = font_cfg.get("name", "Arial")
    font_size = float(font_cfg.get("size_pt", 12))

    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(font_size)

    # Estilos Heading
    for level, size_pt, is_bold in [
        (1, 14, True), (2, 12, True), (3, 12, False), (4, 12, True), (5, 12, False),
    ]:
        style_name = f"Heading {level}"
        if style_name in doc.styles:
            st = doc.styles[style_name]
            st.font.name = font_name
            st.font.size = Pt(size_pt)
            st.font.bold = is_bold

def add_page_numbers(doc: Document, font_name: str = "Arial", font_size_pt: float = 10):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if p.runs:
            for r in p.runs: r.text = ""
        run = p.add_run()
        run.font.name = font_name
        run.font.size = Pt(font_size_pt)
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        run._r.append(fld)

def add_center_line(doc: Document, text: str, size=12, bold=False, uppercase=False, spacing_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper() if uppercase else text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(spacing_after)
    return p

def add_heading(doc: Document, text: str, level: int = 1, spacing_after: int = 0):
    p = doc.add_paragraph(text)
    p.style = f"Heading {level}"
    if p.runs:
        p.runs[0].font.name = "Arial"
        p.runs[0].font.size = Pt(12)
    p.paragraph_format.space_after = Pt(spacing_after)
    return p

def add_page_blocks(doc: Document, blocks: list, default_title_level: int = 4):
    for blk in blocks:
        title = blk.get("title", "")
        lvl = int(blk.get("title_level", default_title_level))
        lines = blk.get("lines", [])
        page_break_after = blk.get("page_break_after", True)

        if title: add_heading(doc, title, level=lvl)
        for line in lines: doc.add_paragraph(line)
        if page_break_after: doc.add_page_break()

def add_center_logo(doc: Document, logo_path: str, width_cm: float = 3.5, spacing_after_pt: int = 6):
    if not logo_path or not os.path.exists(logo_path):
        print(f"[WARN] Logo no encontrado en: {logo_path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(logo_path, width=Cm(width_cm))
    p.paragraph_format.space_after = Pt(spacing_after_pt)

# -------------------------
# WORD: CAMPOS (INDICES)
# -------------------------

def add_field(paragraph, instr: str):
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr)
    run._r.append(fld)

def add_toc_page(doc: Document, toc_cfg: dict):
    p = doc.add_paragraph("INDICE")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0]; r.bold = True; r.font.name = "Arial"; r.font.size = Pt(12)

    min_lv = int(toc_cfg.get("min_level", 1))
    max_lv = int(toc_cfg.get("max_level", 3))
    toc_p = doc.add_paragraph()
    add_field(toc_p, f'TOC \\\\o "{min_lv}-{max_lv}" \\\\h \\\\z \\\\u')
    doc.add_page_break()

def add_list_of_tables(doc: Document):
    p = doc.add_paragraph("INDICE DE TABLAS")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0]; r.bold = True; r.font.name = "Arial"; r.font.size = Pt(12)
    lot_p = doc.add_paragraph()
    add_field(lot_p, 'TOC \\\\h \\\\z \\\\c "Tabla"')
    doc.add_page_break()

def add_list_of_figures(doc: Document):
    p = doc.add_paragraph("INDICE DE FIGURAS")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0]; r.bold = True; r.font.name = "Arial"; r.font.size = Pt(12)
    lof_p = doc.add_paragraph()
    add_field(lof_p, 'TOC \\\\h \\\\z \\\\c "Figura"')
    doc.add_page_break()

# -------------------------
# CARATULA DESDE JSON
# -------------------------

def add_cover_from_cfg(doc: Document, cfg: dict, base_dir: str):
    cover = cfg["cover"]
    # Buscamos el logo relativo al script
    logo_path = resolve_path(base_dir, cfg.get("logo_path", ""))
    
    add_center_logo(doc, logo_path, width_cm=float(cover.get("logo_width_cm", 3.5)), spacing_after_pt=6)

    text_size = int(cover.get("text_size_pt", 12))
    title_size = int(cover.get("title_size_pt", 14))

    add_center_line(doc, cover["universidad_linea"], size=text_size, bold=True, uppercase=True, spacing_after=6)
    add_center_line(doc, cover["unidad"], size=text_size, bold=False, uppercase=True, spacing_after=18)
    add_center_line(doc, f'"{cover["titulo"]}"', size=title_size, bold=True, uppercase=True, spacing_after=12)
    add_center_line(doc, f"TESIS PARA OPTAR EL GRADO ACADEMICO DE {cover['grado_maestria']}", size=text_size, bold=True, uppercase=True, spacing_after=18)

    add_center_line(doc, cover["autor"], size=text_size, uppercase=True, spacing_after=6)
    add_center_line(doc, cover["asesor"], size=text_size, uppercase=True, spacing_after=6)
    add_center_line(doc, f"LINEA DE INVESTIGACION: {cover['linea']}", size=text_size, uppercase=True, spacing_after=18)

    add_center_line(doc, f"{cover.get('ciudad','')}, {cover['anio']}", size=text_size, uppercase=True, spacing_after=6)
    add_center_line(doc, cover.get("pais", "PERU"), size=text_size, uppercase=True, spacing_after=0)
    doc.add_page_break()

# -------------------------
# ESTRUCTURA DESDE JSON
# -------------------------

def add_structure_from_cfg(doc: Document, cfg: dict):
    rules = cfg.get("structure_rules", {})
    add_placeholder = bool(rules.get("add_placeholder_after_heading", True))
    break_after_level1 = bool(rules.get("page_break_after_level_1", True))
    structure = cfg.get("structure", [])
    total = len(structure)

    for i, item in enumerate(structure):
        lvl = int(item["level"])
        title = item["title"]
        add_heading(doc, title, level=lvl)

        if add_placeholder and bool(item.get("placeholder", True)):
            doc.add_paragraph("{{COMPLETAR}}")

        extra_lines = item.get("lines", [])
        if extra_lines:
            for line in extra_lines: doc.add_paragraph(str(line))

        if break_after_level1 and lvl == 1 and i < total - 1:
            next_lvl = int(structure[i + 1]["level"])
            if next_lvl == 1: doc.add_page_break()

# -------------------------
# MAIN GENERATOR
# -------------------------

def generate(config_path: str, output_path_override: str = None):
    # La carpeta base es donde esta este script
    base_dir = os.path.dirname(os.path.abspath(__file__))
    
    # 1. Cargar configuración
    if not os.path.exists(config_path):
        # Intento de buscar en formats/ si solo nos pasaron el nombre
        possible_path = os.path.join(base_dir, "formats", os.path.basename(config_path))
        if os.path.exists(possible_path):
            config_path = possible_path
        else:
            raise FileNotFoundError(f"Config JSON no encontrado en: {config_path}")

    cfg = load_json(config_path)

    # 2. Crear Documento
    doc = Document()
    set_page_setup(doc, cfg)
    add_cover_from_cfg(doc, cfg, base_dir)

    pre_pages = cfg.get("pre_pages", [])
    if pre_pages: add_page_blocks(doc, pre_pages, default_title_level=4)

    add_toc_page(doc, cfg.get("toc", {"min_level": 1, "max_level": 3}))
    if cfg.get("include_list_of_tables", False): add_list_of_tables(doc)
    if cfg.get("include_list_of_figures", False): add_list_of_figures(doc)

    add_structure_from_cfg(doc, cfg)
    add_page_numbers(doc)

    # 3. Guardar
    if output_path_override:
        # MODO SERVIDOR: Usamos la ruta que nos da el server (carpeta descargas)
        final_path = output_path_override
    else:
        # MODO MANUAL: Guardamos en la carpeta del script
        output_name = cfg.get("output_name", "output.docx")
        final_path = os.path.join(base_dir, output_name)

    doc.save(final_path)
    
    # IMPORTANTE: Usamos [OK] en lugar de emojis para evitar crash en Windows
    print(f"[OK] Documento guardado en: {final_path}")

    # Si NO estamos en modo servidor (output_override es None), abrimos el archivo
    if not output_path_override:
        open_document(final_path)

if __name__ == "__main__":
    
    # ----------------------------------------------------
    # MODO SERVIDOR (AUTOMÁTICO)
    # Recibe: script.py [json_path] [output_path]
    # ----------------------------------------------------
    if len(sys.argv) > 2:
        json_arg = sys.argv[1]
        output_arg = sys.argv[2]
        
        try:
            generate(json_arg, output_path_override=output_arg)
        except Exception as e:
            print(f"[ERROR] Fallo critico: {e}")
            sys.exit(1)

    # ----------------------------------------------------
    # MODO MANUAL (CLI)
    # ----------------------------------------------------
    else:
        base_dir = os.path.dirname(os.path.abspath(__file__))
        print("="*40)
        print("   GENERADOR MAESTRIA UNAC (CLI)")
        print("="*40)
        print("1. Maestría Cualitativa")
        print("2. Maestría Cuantitativa")
        
        try: op = input(">> Opcion (1/2): ").strip()
        except: sys.exit()

        json_file = ""
        if op == "1": json_file = "unac_maestria_cual.json"
        elif op == "2": json_file = "unac_maestria_cuant.json"
        else: print("Opcion no valida"); sys.exit()

        config_path = os.path.join(base_dir, "formats", json_file)
        generate(config_path)