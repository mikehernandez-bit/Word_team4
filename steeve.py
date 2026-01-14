# IMPORTANTE: Antes de ejecutar este script, instala la biblioteca:
# pip install python-docx

import os
import subprocess
import platform

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("ERROR: La biblioteca 'python-docx' no está instalada.")
    print("Por favor, ejecuta en tu terminal:")
    print("  pip install python-docx")
    print("\nO si usas Python 3:")
    print("  pip3 install python-docx")
    exit()

# Crear un nuevo documento
doc = Document()

# Agregar un título
titulo = doc.add_heading('Documento de Ejemplo', 0)
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Agregar un subtítulo
doc.add_heading('Introducción', level=1)

# Agregar un párrafo
parrafo1 = doc.add_paragraph(
    'Este es un documento Word simple creado con Python. '
    'La biblioteca python-docx permite crear y modificar documentos '
    'de Microsoft Word de manera programática.'
)

# Agregar otro subtítulo
doc.add_heading('Características', level=1)

# Agregar una lista con viñetas
doc.add_paragraph('Fácil de usar', style='List Bullet')
doc.add_paragraph('Soporta estilos y formato', style='List Bullet')
doc.add_paragraph('Permite insertar tablas e imágenes', style='List Bullet')

# Agregar una tabla simple
doc.add_heading('Tabla de Ejemplo', level=1)
tabla = doc.add_table(rows=3, cols=3)
tabla.style = 'Light Grid Accent 1'

# Llenar la tabla
encabezados = ['Columna 1', 'Columna 2', 'Columna 3']
for i, encabezado in enumerate(encabezados):
    tabla.rows[0].cells[i].text = encabezado

for fila in range(1, 3):
    for col in range(3):
        tabla.rows[fila].cells[col].text = f'Fila {fila}, Col {col+1}'

# Agregar un párrafo final
doc.add_heading('Conclusión', level=1)
doc.add_paragraph(
    'Este es un ejemplo básico. En el futuro podremos adaptarlo '
    'para crear documentos con formato de tesis.'
)

# Guardar el documento
# Puedes cambiar el nombre o la ruta aquí:
nombre_archivo = 'documento_simple.docx'
# O usar ruta completa: 'C:/Users/Steeve/Desktop/mi_documento.docx'

doc.save(nombre_archivo)
print(f"✓ Documento '{nombre_archivo}' creado exitosamente")
print(f"✓ Ubicación: {os.path.abspath(nombre_archivo)}")

# Abrir automáticamente el documento
try:
    sistema = platform.system()
    ruta_completa = os.path.abspath(nombre_archivo)
    
    if sistema == 'Windows':
        os.startfile(ruta_completa)
    elif sistema == 'Darwin':  # macOS
        subprocess.run(['open', ruta_completa])
    else:  # Linux
        subprocess.run(['xdg-open', ruta_completa])
    
    print(f"✓ Abriendo documento automáticamente...")
except Exception as e:
    print(f"⚠ No se pudo abrir automáticamente: {e}")
    print(f"  Puedes abrirlo manualmente desde: {os.path.abspath(nombre_archivo)}")