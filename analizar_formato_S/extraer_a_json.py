"""
Extractor Completo - Convierte documento Word a JSON detallado
Extrae TODOS los elementos: párrafos, tablas, imágenes, estilos, formato, etc.
"""

import os
import json
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import base64

# Colores para terminal
class Color:
    GREEN = '\033[92m'
    YELLOW = '\033[93m'
    RED = '\033[91m'
    BLUE = '\033[94m'
    CYAN = '\033[96m'
    MAGENTA = '\033[95m'
    END = '\033[0m'

class ExtractorCompleto:
    def __init__(self, ruta_documento):
        self.ruta = ruta_documento
        self.doc = None
        self.nombre_json = "estructura_completa.json"  # Nombre por defecto
        self.estructura = {
            'metadata': {},
            'configuracion_pagina': {},
            'secciones': [],
            'elementos': [],
            'encabezados': [],
            'pies_pagina': [],
            'estilos_usados': {}
        }
    
    def cargar_documento(self):
        """Carga el documento Word"""
        try:
            self.doc = Document(self.ruta)
            print(f"{Color.GREEN}✓{Color.END} Documento cargado: {self.ruta}")
            return True
        except Exception as e:
            print(f"{Color.RED}✗{Color.END} Error al cargar documento: {e}")
            return False
    
    def extraer_metadata(self):
        """Extrae metadata del documento"""
        print(f"\n{Color.CYAN}📊 Extrayendo metadata...{Color.END}")
        
        props = self.doc.core_properties
        
        self.estructura['metadata'] = {
            'autor': props.author or '',
            'titulo': props.title or '',
            'asunto': props.subject or '',
            'creado': str(props.created) if props.created else '',
            'modificado': str(props.modified) if props.modified else '',
        }
        
        print(f"  · Título: {Color.BLUE}{self.estructura['metadata']['titulo'] or 'Sin título'}{Color.END}")
    
    def extraer_configuracion_pagina(self):
        """Extrae configuración de página de todas las secciones"""
        print(f"\n{Color.CYAN}📏 Extrayendo configuración de página...{Color.END}")
        
        for i, section in enumerate(self.doc.sections):
            config = {
                'seccion_num': i,
                'ancho_cm': round(section.page_width.cm, 2),
                'alto_cm': round(section.page_height.cm, 2),
                'margen_superior_cm': round(section.top_margin.cm, 2),
                'margen_inferior_cm': round(section.bottom_margin.cm, 2),
                'margen_izquierdo_cm': round(section.left_margin.cm, 2),
                'margen_derecho_cm': round(section.right_margin.cm, 2),
                'orientacion': 'horizontal' if section.orientation == 1 else 'vertical',
            }
            
            self.estructura['secciones'].append(config)
        
        # Guardar la configuración principal (primera sección)
        if self.estructura['secciones']:
            self.estructura['configuracion_pagina'] = self.estructura['secciones'][0]
        
        print(f"  · Secciones encontradas: {Color.BLUE}{len(self.estructura['secciones'])}{Color.END}")
        print(f"  · Tamaño página: {Color.BLUE}{self.estructura['configuracion_pagina']['ancho_cm']} x {self.estructura['configuracion_pagina']['alto_cm']} cm{Color.END}")
    
    def extraer_encabezados_pies(self):
        """Extrae encabezados y pies de página de todas las secciones"""
        print(f"\n{Color.CYAN}📋 Extrayendo encabezados y pies de página...{Color.END}")
        
        for i, section in enumerate(self.doc.sections):
            # Encabezado
            header_elementos = []
            for p in section.header.paragraphs:
                header_elementos.append(self._extraer_parrafo_completo(p))
            
            if any(e['texto'].strip() for e in header_elementos):
                self.estructura['encabezados'].append({
                    'seccion': i,
                    'elementos': header_elementos
                })
            
            # Pie de página
            footer_elementos = []
            for p in section.footer.paragraphs:
                footer_elementos.append(self._extraer_parrafo_completo(p))
            
            if any(e['texto'].strip() for e in footer_elementos):
                self.estructura['pies_pagina'].append({
                    'seccion': i,
                    'elementos': footer_elementos
                })
        
        print(f"  · Encabezados: {Color.BLUE}{len(self.estructura['encabezados'])}{Color.END}")
        print(f"  · Pies de página: {Color.BLUE}{len(self.estructura['pies_pagina'])}{Color.END}")
    
    def _obtener_alineacion(self, alignment):
        """Convierte código de alineación a texto"""
        alineaciones = {
            None: 'left',
            0: 'left',
            1: 'center',
            2: 'right',
            3: 'justify'
        }
        return alineaciones.get(alignment, 'left')
    
    def _extraer_formato_run(self, run):
        """Extrae formato completo de un run de texto"""
        formato = {
            'texto': run.text,
            'negrita': run.bold if run.bold is not None else False,
            'cursiva': run.italic if run.italic is not None else False,
            'subrayado': run.underline if run.underline is not None else False,
            'tachado': run.font.strike if run.font.strike is not None else False,
        }
        
        # Fuente
        if run.font.name:
            formato['fuente'] = run.font.name
        
        # Tamaño
        if run.font.size:
            formato['tamaño_pt'] = run.font.size.pt
        
        # Color
        if run.font.color and run.font.color.rgb:
            formato['color_rgb'] = str(run.font.color.rgb)
        
        # Resaltado
        if run.font.highlight_color:
            formato['resaltado'] = str(run.font.highlight_color)
        
        return formato
    
    def _extraer_parrafo_completo(self, parrafo):
        """Extrae toda la información de un párrafo"""
        info_parrafo = {
            'tipo': 'parrafo',
            'texto': parrafo.text,
            'estilo': parrafo.style.name if parrafo.style else 'Normal',
            'alineacion': self._obtener_alineacion(parrafo.alignment),
            'runs': []
        }
        
        # Formato de párrafo
        fmt = parrafo.paragraph_format
        
        if fmt.left_indent:
            info_parrafo['sangria_izquierda_cm'] = round(fmt.left_indent.cm, 2)
        
        if fmt.right_indent:
            info_parrafo['sangria_derecha_cm'] = round(fmt.right_indent.cm, 2)
        
        if fmt.first_line_indent:
            info_parrafo['primera_linea_cm'] = round(fmt.first_line_indent.cm, 2)
        
        if fmt.space_before:
            info_parrafo['espacio_antes_pt'] = fmt.space_before.pt
        
        if fmt.space_after:
            info_parrafo['espacio_despues_pt'] = fmt.space_after.pt
        
        if fmt.line_spacing:
            info_parrafo['interlineado'] = fmt.line_spacing
        
        # Extraer runs (fragmentos de texto con formato)
        for run in parrafo.runs:
            info_parrafo['runs'].append(self._extraer_formato_run(run))
        
        return info_parrafo
    
    def _extraer_tabla(self, tabla):
        """Extrae información completa de una tabla"""
        info_tabla = {
            'tipo': 'tabla',
            'filas': len(tabla.rows),
            'columnas': len(tabla.columns),
            'estilo': tabla.style.name if tabla.style else 'Sin estilo',
            'contenido': []
        }
        
        # Extraer contenido de cada celda
        for i, row in enumerate(tabla.rows):
            fila_data = []
            for j, cell in enumerate(row.cells):
                celda_info = {
                    'texto': cell.text,
                    'parrafos': []
                }
                
                # Extraer párrafos dentro de la celda
                for p in cell.paragraphs:
                    celda_info['parrafos'].append(self._extraer_parrafo_completo(p))
                
                fila_data.append(celda_info)
            
            info_tabla['contenido'].append(fila_data)
        
        return info_tabla
    
    def extraer_elementos_documento(self):
        """Extrae todos los elementos del documento en orden"""
        print(f"\n{Color.CYAN}📄 Extrayendo elementos del documento...{Color.END}")
        
        elementos_procesados = 0
        
        # Iterar sobre todos los elementos del cuerpo del documento
        for element in self.doc.element.body:
            # Párrafos
            if element.tag.endswith('p'):
                # Buscar el párrafo correspondiente
                for parrafo in self.doc.paragraphs:
                    if parrafo._element == element:
                        info_parrafo = self._extraer_parrafo_completo(parrafo)
                        self.estructura['elementos'].append(info_parrafo)
                        elementos_procesados += 1
                        
                        # Registrar estilo usado
                        estilo_nombre = info_parrafo['estilo']
                        if estilo_nombre not in self.estructura['estilos_usados']:
                            self.estructura['estilos_usados'][estilo_nombre] = 1
                        else:
                            self.estructura['estilos_usados'][estilo_nombre] += 1
                        break
            
            # Tablas
            elif element.tag.endswith('tbl'):
                # Buscar la tabla correspondiente
                for tabla in self.doc.tables:
                    if tabla._element == element:
                        info_tabla = self._extraer_tabla(tabla)
                        self.estructura['elementos'].append(info_tabla)
                        elementos_procesados += 1
                        break
        
        print(f"  · Elementos extraídos: {Color.BLUE}{elementos_procesados}{Color.END}")
        print(f"  · Estilos únicos: {Color.BLUE}{len(self.estructura['estilos_usados'])}{Color.END}")
        
        # Mostrar preview de estilos
        for estilo, count in list(self.estructura['estilos_usados'].items())[:5]:
            print(f"    - {estilo}: {count} usos")
    
    def guardar_json(self):
        """Guarda la estructura extraída en JSON"""
        ruta_salida = os.path.join('plantillas', self.nombre_json)
        
        try:
            with open(ruta_salida, 'w', encoding='utf-8') as f:
                json.dump(self.estructura, f, indent=2, ensure_ascii=False)
            
            # Calcular tamaño
            tamaño_kb = os.path.getsize(ruta_salida) / 1024
            
            print(f"\n{Color.GREEN}✓{Color.END} Estructura guardada en: {Color.BLUE}{ruta_salida}{Color.END}")
            print(f"{Color.BLUE}ℹ{Color.END} Tamaño del JSON: {tamaño_kb:.2f} KB")
            
            return ruta_salida
        except Exception as e:
            print(f"\n{Color.RED}✗{Color.END} Error al guardar JSON: {e}")
            return None
    
    def generar_reporte(self):
        """Genera reporte visual de la extracción"""
        print(f"\n{'='*70}")
        print(f"{Color.MAGENTA}REPORTE DE EXTRACCIÓN{Color.END}")
        print(f"{'='*70}\n")
        
        print(f"{Color.YELLOW}📋 METADATA:{Color.END}")
        meta = self.estructura['metadata']
        print(f"  · Título: {meta['titulo'] or 'Sin título'}")
        print(f"  · Autor: {meta['autor'] or 'Sin autor'}")
        
        print(f"\n{Color.YELLOW}📏 CONFIGURACIÓN:{Color.END}")
        config = self.estructura['configuracion_pagina']
        print(f"  · Tamaño: {config['ancho_cm']} x {config['alto_cm']} cm")
        print(f"  · Márgenes: S={config['margen_superior_cm']}, I={config['margen_inferior_cm']}, L={config['margen_izquierdo_cm']}, R={config['margen_derecho_cm']} cm")
        print(f"  · Orientación: {config['orientacion']}")
        
        print(f"\n{Color.YELLOW}📄 CONTENIDO:{Color.END}")
        print(f"  · Total elementos: {len(self.estructura['elementos'])}")
        
        # Contar tipos
        parrafos = sum(1 for e in self.estructura['elementos'] if e['tipo'] == 'parrafo')
        tablas = sum(1 for e in self.estructura['elementos'] if e['tipo'] == 'tabla')
        
        print(f"  · Párrafos: {parrafos}")
        print(f"  · Tablas: {tablas}")
        print(f"  · Estilos únicos: {len(self.estructura['estilos_usados'])}")
        
        print(f"\n{Color.YELLOW}📋 ENCABEZADOS Y PIES:{Color.END}")
        print(f"  · Encabezados: {len(self.estructura['encabezados'])}")
        print(f"  · Pies de página: {len(self.estructura['pies_pagina'])}")
        
        print(f"\n{'='*70}\n")
    
    def ejecutar_extraccion_completa(self):
        """Ejecuta la extracción completa"""
        print(f"\n{Color.BLUE}🔍 INICIANDO EXTRACCIÓN COMPLETA A JSON...{Color.END}\n")
        
        if not self.cargar_documento():
            return False
        
        self.extraer_metadata()
        self.extraer_configuracion_pagina()
        self.extraer_encabezados_pies()
        self.extraer_elementos_documento()
        
        self.generar_reporte()
        
        ruta_json = self.guardar_json()
        
        if ruta_json:
            print(f"\n{Color.GREEN}✅ EXTRACCIÓN COMPLETADA{Color.END}")
            print(f"\nPróximo paso: Ejecuta {Color.BLUE}replicar_desde_json.py{Color.END} para recrear el documento\n")
            return True
        
        return False

def listar_documentos_plantilla():
    """Lista todos los documentos .docx en la carpeta plantillas"""
    carpeta_plantillas = 'plantillas'
    
    if not os.path.exists(carpeta_plantillas):
        print(f"{Color.RED}✗{Color.END} La carpeta 'plantillas' no existe")
        print(f"{Color.YELLOW}⚠{Color.END} Ejecuta primero: {Color.BLUE}config_plantilla.py{Color.END}")
        return None
    
    # Buscar todos los archivos .docx
    documentos = [f for f in os.listdir(carpeta_plantillas) 
                  if f.endswith('.docx') or f.endswith('.doc')]
    
    if not documentos:
        print(f"{Color.RED}✗{Color.END} No se encontraron documentos Word en la carpeta 'plantillas'")
        print(f"{Color.YELLOW}⚠{Color.END} Ejecuta primero: {Color.BLUE}config_plantilla.py{Color.END}")
        return None
    
    return documentos

def seleccionar_documento(documentos):
    """Permite al usuario seleccionar un documento de la lista"""
    print(f"\n{Color.CYAN}📂 Documentos disponibles en 'plantillas':{Color.END}\n")
    
    for i, doc in enumerate(documentos, 1):
        tamaño_kb = os.path.getsize(os.path.join('plantillas', doc)) / 1024
        print(f"  {Color.BLUE}[{i}]{Color.END} {doc} ({tamaño_kb:.2f} KB)")
    
    print()
    
    while True:
        try:
            seleccion = input(f"Selecciona un documento (1-{len(documentos)}): ").strip()
            
            if not seleccion:
                print(f"{Color.YELLOW}⚠{Color.END} Debes seleccionar un documento")
                continue
            
            idx = int(seleccion) - 1
            
            if 0 <= idx < len(documentos):
                return documentos[idx]
            else:
                print(f"{Color.RED}✗{Color.END} Número fuera de rango. Intenta de nuevo.")
        except ValueError:
            print(f"{Color.RED}✗{Color.END} Por favor ingresa un número válido")
        except KeyboardInterrupt:
            print(f"\n{Color.YELLOW}⚠{Color.END} Operación cancelada")
            return None

def main():
    print("="*70)
    print(f"{Color.BLUE}EXTRACTOR COMPLETO - WORD A JSON{Color.END}")
    print("="*70)
    
    # Listar documentos disponibles
    documentos = listar_documentos_plantilla()
    
    if not documentos:
        return
    
    # Si solo hay un documento, usarlo automáticamente
    if len(documentos) == 1:
        documento_seleccionado = documentos[0]
        print(f"\n{Color.GREEN}✓{Color.END} Usando documento: {Color.BLUE}{documento_seleccionado}{Color.END}")
    else:
        # Permitir selección
        documento_seleccionado = seleccionar_documento(documentos)
        
        if not documento_seleccionado:
            return
    
    # Ruta completa del documento
    ruta_plantilla = os.path.join('plantillas', documento_seleccionado)
    
    # Preguntar nombre del JSON de salida
    print(f"\n{Color.CYAN}💾 Nombre del archivo JSON a generar:{Color.END}")
    print(f"{Color.YELLOW}Ejemplo:{Color.END} mi_formato, estructura_tesis, etc.")
    print(f"{Color.YELLOW}Dejar vacío para usar:{Color.END} estructura_completa.json")
    
    nombre_json = input("\nNombre (sin extensión): ").strip()
    
    if not nombre_json:
        nombre_json = "estructura_completa"
    
    # Limpiar nombre (quitar caracteres no válidos)
    nombre_json = nombre_json.replace(' ', '_').replace('.json', '')
    nombre_json = f"{nombre_json}.json"
    
    print()
    
    # Ejecutar extracción
    extractor = ExtractorCompleto(ruta_plantilla)
    extractor.nombre_json = nombre_json  # Pasar nombre personalizado
    extractor.ejecutar_extraccion_completa()

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print(f"\n\n{Color.YELLOW}⚠{Color.END} Extracción cancelada.")
    except Exception as e:
        print(f"\n{Color.RED}✗{Color.END} Error inesperado: {e}")
        import traceback
        traceback.print_exc()