"""
Replicador de Documentos - Recrea documento Word desde JSON
Replica EXACTAMENTE la estructura, formato y elementos del documento original
"""

import os
import json
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess
import platform

# Colores para terminal
class Color:
    GREEN = '\033[92m'
    YELLOW = '\033[93m'
    RED = '\033[91m'
    BLUE = '\033[94m'
    CYAN = '\033[96m'
    MAGENTA = '\033[95m'
    END = '\033[0m'

class ReplicadorDocumento:
    def __init__(self):
        self.json_path = None  # Se seleccionará después
        self.estructura = None
        self.doc = None
        self.nombre_archivo = None
    
    def cargar_json(self):
        """Carga la estructura desde JSON"""
        if not self.json_path or not os.path.exists(self.json_path):
            print(f"{Color.RED}✗{Color.END} No se encontró el JSON en: {self.json_path}")
            print(f"{Color.YELLOW}⚠{Color.END} Ejecuta primero: {Color.BLUE}extraer_a_json.py{Color.END}")
            return False
        
        try:
            with open(self.json_path, 'r', encoding='utf-8') as f:
                self.estructura = json.load(f)
            
            print(f"\n{Color.GREEN}✓{Color.END} Estructura JSON cargada exitosamente")
            print(f"{Color.BLUE}ℹ{Color.END} Elementos a replicar: {len(self.estructura['elementos'])}")
            
            # Mostrar preview
            meta = self.estructura.get('metadata', {})
            if meta.get('titulo'):
                print(f"{Color.BLUE}ℹ{Color.END} Título del documento: {meta['titulo']}")
            
            return True
        except Exception as e:
            print(f"{Color.RED}✗{Color.END} Error al cargar JSON: {e}")
            return False
    
    def crear_documento(self):
        """Crea un nuevo documento Word"""
        try:
            self.doc = Document()
            print(f"{Color.GREEN}✓{Color.END} Nuevo documento creado")
            return True
        except Exception as e:
            print(f"{Color.RED}✗{Color.END} Error al crear documento: {e}")
            return False
    
    def aplicar_configuracion_pagina(self):
        """Aplica la configuración de página del JSON"""
        print(f"\n{Color.CYAN}📏 Aplicando configuración de página...{Color.END}")
        
        try:
            config = self.estructura['configuracion_pagina']
            
            for section in self.doc.sections:
                section.page_width = Cm(config['ancho_cm'])
                section.page_height = Cm(config['alto_cm'])
                section.top_margin = Cm(config['margen_superior_cm'])
                section.bottom_margin = Cm(config['margen_inferior_cm'])
                section.left_margin = Cm(config['margen_izquierdo_cm'])
                section.right_margin = Cm(config['margen_derecho_cm'])
                
                if config['orientacion'] == 'horizontal':
                    section.orientation = 1
                else:
                    section.orientation = 0
            
            print(f"  {Color.GREEN}✓{Color.END} Tamaño: {config['ancho_cm']} x {config['alto_cm']} cm")
            print(f"  {Color.GREEN}✓{Color.END} Márgenes aplicados")
            return True
        except Exception as e:
            print(f"  {Color.YELLOW}⚠{Color.END} Advertencia: {e}")
            return False
    
    def aplicar_encabezados_pies(self):
        """Aplica encabezados y pies de página"""
        print(f"\n{Color.CYAN}📋 Aplicando encabezados y pies...{Color.END}")
        
        try:
            # Encabezados
            for enc in self.estructura['encabezados']:
                seccion_num = enc['seccion']
                if seccion_num < len(self.doc.sections):
                    section = self.doc.sections[seccion_num]
                    
                    for elem in enc['elementos']:
                        p = section.header.add_paragraph()
                        self._aplicar_formato_parrafo(p, elem)
            
            # Pies de página
            for pie in self.estructura['pies_pagina']:
                seccion_num = pie['seccion']
                if seccion_num < len(self.doc.sections):
                    section = self.doc.sections[seccion_num]
                    
                    for elem in pie['elementos']:
                        p = section.footer.add_paragraph()
                        self._aplicar_formato_parrafo(p, elem)
            
            print(f"  {Color.GREEN}✓{Color.END} Encabezados: {len(self.estructura['encabezados'])}")
            print(f"  {Color.GREEN}✓{Color.END} Pies: {len(self.estructura['pies_pagina'])}")
            return True
        except Exception as e:
            print(f"  {Color.YELLOW}⚠{Color.END} Advertencia: {e}")
            return False
    
    def _obtener_alineacion_enum(self, alineacion_texto):
        """Convierte texto de alineación a enum"""
        alineaciones = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        return alineaciones.get(alineacion_texto, WD_ALIGN_PARAGRAPH.LEFT)
    
    def _aplicar_formato_run(self, run, formato_run):
        """Aplica formato a un run de texto"""
        # Texto
        run.text = formato_run['texto']
        
        # Negrita, cursiva, subrayado
        if formato_run.get('negrita'):
            run.bold = True
        
        if formato_run.get('cursiva'):
            run.italic = True
        
        if formato_run.get('subrayado'):
            run.underline = True
        
        if formato_run.get('tachado'):
            run.font.strike = True
        
        # Fuente
        if 'fuente' in formato_run:
            run.font.name = formato_run['fuente']
        
        # Tamaño
        if 'tamaño_pt' in formato_run:
            run.font.size = Pt(formato_run['tamaño_pt'])
        
        # Color
        if 'color_rgb' in formato_run:
            try:
                # Convertir string RGB a objeto RGBColor
                color_str = formato_run['color_rgb'].replace('RGBColor(0x', '').replace(')', '')
                if len(color_str) == 6:
                    r = int(color_str[0:2], 16)
                    g = int(color_str[2:4], 16)
                    b = int(color_str[4:6], 16)
                    run.font.color.rgb = RGBColor(r, g, b)
            except:
                pass
    
    def _aplicar_formato_parrafo(self, parrafo, info_parrafo):
        """Aplica formato completo a un párrafo"""
        # Alineación
        parrafo.alignment = self._obtener_alineacion_enum(info_parrafo['alineacion'])
        
        # Formato de párrafo
        fmt = parrafo.paragraph_format
        
        if 'sangria_izquierda_cm' in info_parrafo:
            fmt.left_indent = Cm(info_parrafo['sangria_izquierda_cm'])
        
        if 'sangria_derecha_cm' in info_parrafo:
            fmt.right_indent = Cm(info_parrafo['sangria_derecha_cm'])
        
        if 'primera_linea_cm' in info_parrafo:
            fmt.first_line_indent = Cm(info_parrafo['primera_linea_cm'])
        
        if 'espacio_antes_pt' in info_parrafo:
            fmt.space_before = Pt(info_parrafo['espacio_antes_pt'])
        
        if 'espacio_despues_pt' in info_parrafo:
            fmt.space_after = Pt(info_parrafo['espacio_despues_pt'])
        
        if 'interlineado' in info_parrafo:
            fmt.line_spacing = info_parrafo['interlineado']
        
        # Aplicar runs (fragmentos de texto con formato)
        if 'runs' in info_parrafo and info_parrafo['runs']:
            # Limpiar párrafo
            parrafo.clear()
            
            for run_info in info_parrafo['runs']:
                run = parrafo.add_run()
                self._aplicar_formato_run(run, run_info)
        else:
            # Si no hay runs, usar el texto completo
            if info_parrafo['texto']:
                parrafo.add_run(info_parrafo['texto'])
    
    def _replicar_tabla(self, info_tabla):
        """Replica una tabla exactamente"""
        try:
            # Crear tabla
            tabla = self.doc.add_table(
                rows=info_tabla['filas'],
                cols=info_tabla['columnas']
            )
            
            # Aplicar estilo si existe
            if info_tabla['estilo'] and info_tabla['estilo'] != 'Sin estilo':
                try:
                    tabla.style = info_tabla['estilo']
                except:
                    pass
            
            # Llenar contenido
            for i, fila_data in enumerate(info_tabla['contenido']):
                for j, celda_info in enumerate(fila_data):
                    celda = tabla.rows[i].cells[j]
                    
                    # Limpiar celda
                    celda._element.clear_content()
                    
                    # Agregar párrafos de la celda
                    if celda_info['parrafos']:
                        for p_info in celda_info['parrafos']:
                            p = celda.add_paragraph()
                            self._aplicar_formato_parrafo(p, p_info)
                    else:
                        # Si no hay párrafos, agregar texto simple
                        if celda_info['texto']:
                            celda.text = celda_info['texto']
            
            return True
        except Exception as e:
            print(f"  {Color.YELLOW}⚠{Color.END} Error al replicar tabla: {e}")
            return False
    
    def replicar_elementos(self):
        """Replica todos los elementos del documento en orden"""
        print(f"\n{Color.CYAN}📄 Replicando elementos del documento...{Color.END}")
        
        elementos_replicados = 0
        
        for i, elemento in enumerate(self.estructura['elementos']):
            try:
                if elemento['tipo'] == 'parrafo':
                    # Crear párrafo
                    p = self.doc.add_paragraph()
                    self._aplicar_formato_parrafo(p, elemento)
                    elementos_replicados += 1
                
                elif elemento['tipo'] == 'tabla':
                    # Crear tabla
                    if self._replicar_tabla(elemento):
                        elementos_replicados += 1
                
                # Mostrar progreso cada 20 elementos
                if (i + 1) % 20 == 0:
                    print(f"  · Procesados: {i + 1}/{len(self.estructura['elementos'])}")
            
            except Exception as e:
                print(f"  {Color.YELLOW}⚠{Color.END} Error en elemento {i}: {e}")
        
        print(f"  {Color.GREEN}✓{Color.END} Elementos replicados: {Color.BLUE}{elementos_replicados}/{len(self.estructura['elementos'])}{Color.END}")
        return elementos_replicados > 0
    
    def guardar_documento(self):
        """Guarda el documento replicado"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        self.nombre_archivo = f'replica_{timestamp}.docx'
        ruta_completa = os.path.join('generados', self.nombre_archivo)
        
        try:
            # Asegurar que existe la carpeta
            os.makedirs('generados', exist_ok=True)
            
            self.doc.save(ruta_completa)
            
            tamaño_kb = os.path.getsize(ruta_completa) / 1024
            
            print(f"\n{Color.GREEN}✓{Color.END} Documento guardado: {Color.BLUE}{ruta_completa}{Color.END}")
            print(f"{Color.BLUE}ℹ{Color.END} Tamaño: {tamaño_kb:.2f} KB")
            
            return ruta_completa
        except Exception as e:
            print(f"\n{Color.RED}✗{Color.END} Error al guardar: {e}")
            return None
    
    def abrir_documento(self, ruta):
        """Abre el documento automáticamente"""
        try:
            sistema = platform.system()
            ruta_abs = os.path.abspath(ruta)
            
            if sistema == 'Windows':
                os.startfile(ruta_abs)
            elif sistema == 'Darwin':
                subprocess.run(['open', ruta_abs])
            else:
                subprocess.run(['xdg-open', ruta_abs])
            
            print(f"{Color.GREEN}✓{Color.END} Abriendo documento...")
            return True
        except Exception as e:
            print(f"{Color.YELLOW}⚠{Color.END} No se pudo abrir: {e}")
            return False
    
    def ejecutar_replicacion(self):
        """Ejecuta el proceso completo de replicación"""
        print(f"\n{Color.MAGENTA}{'='*70}{Color.END}")
        print(f"{Color.MAGENTA}REPLICANDO DOCUMENTO DESDE JSON{Color.END}")
        print(f"{Color.MAGENTA}{'='*70}{Color.END}")
        
        if not self.cargar_json():
            return False
        
        if not self.crear_documento():
            return False
        
        self.aplicar_configuracion_pagina()
        self.aplicar_encabezados_pies()
        
        if not self.replicar_elementos():
            print(f"\n{Color.RED}✗{Color.END} No se pudo replicar el contenido")
            return False
        
        ruta_guardada = self.guardar_documento()
        
        if ruta_guardada:
            print(f"\n{'='*70}")
            print(f"{Color.GREEN}✅ REPLICACIÓN COMPLETADA{Color.END}")
            print(f"{'='*70}\n")
            
            # Preguntar si abrir
            respuesta = input(f"¿Abrir el documento ahora? (s/n): ").lower().strip()
            if respuesta == 's':
                self.abrir_documento(ruta_guardada)
            
            print()
            return True
        
        return False

def listar_archivos_json():
    """Lista todos los archivos JSON en la carpeta plantillas"""
    carpeta_plantillas = 'plantillas'
    
    if not os.path.exists(carpeta_plantillas):
        print(f"{Color.RED}✗{Color.END} La carpeta 'plantillas' no existe")
        return None
    
    # Buscar todos los archivos .json
    archivos_json = [f for f in os.listdir(carpeta_plantillas) if f.endswith('.json')]
    
    if not archivos_json:
        print(f"{Color.RED}✗{Color.END} No se encontraron archivos JSON en 'plantillas'")
        print(f"{Color.YELLOW}⚠{Color.END} Ejecuta primero: {Color.BLUE}extraer_a_json.py{Color.END}")
        return None
    
    return archivos_json

def seleccionar_json(archivos_json):
    """Permite al usuario seleccionar un archivo JSON"""
    print(f"\n{Color.CYAN}📂 Archivos JSON disponibles en 'plantillas':{Color.END}\n")
    
    for i, archivo in enumerate(archivos_json, 1):
        ruta_completa = os.path.join('plantillas', archivo)
        tamaño_kb = os.path.getsize(ruta_completa) / 1024
        
        # Intentar leer metadata del JSON
        try:
            with open(ruta_completa, 'r', encoding='utf-8') as f:
                data = json.load(f)
                titulo = data.get('metadata', {}).get('titulo', 'Sin título')
                elementos = len(data.get('elementos', []))
                print(f"  {Color.BLUE}[{i}]{Color.END} {archivo}")
                print(f"      └─ Título: {titulo}")
                print(f"      └─ Elementos: {elementos}")
                print(f"      └─ Tamaño: {tamaño_kb:.2f} KB")
        except:
            print(f"  {Color.BLUE}[{i}]{Color.END} {archivo} ({tamaño_kb:.2f} KB)")
        
        print()
    
    while True:
        try:
            seleccion = input(f"Selecciona un archivo JSON (1-{len(archivos_json)}): ").strip()
            
            if not seleccion:
                print(f"{Color.YELLOW}⚠{Color.END} Debes seleccionar un archivo")
                continue
            
            idx = int(seleccion) - 1
            
            if 0 <= idx < len(archivos_json):
                return archivos_json[idx]
            else:
                print(f"{Color.RED}✗{Color.END} Número fuera de rango. Intenta de nuevo.")
        except ValueError:
            print(f"{Color.RED}✗{Color.END} Por favor ingresa un número válido")
        except KeyboardInterrupt:
            print(f"\n{Color.YELLOW}⚠{Color.END} Operación cancelada")
            return None

def main():
    print("="*70)
    print(f"{Color.BLUE}REPLICADOR DE DOCUMENTOS DESDE JSON{Color.END}")
    print("="*70)
    print()
    
    # Listar archivos JSON disponibles
    archivos_json = listar_archivos_json()
    
    if not archivos_json:
        return
    
    # Si solo hay un JSON, usarlo automáticamente
    if len(archivos_json) == 1:
        json_seleccionado = archivos_json[0]
        print(f"{Color.GREEN}✓{Color.END} Usando JSON: {Color.BLUE}{json_seleccionado}{Color.END}")
    else:
        # Permitir selección
        json_seleccionado = seleccionar_json(archivos_json)
        
        if not json_seleccionado:
            return
    
    # Crear replicador con el JSON seleccionado
    replicador = ReplicadorDocumento()
    replicador.json_path = os.path.join('plantillas', json_seleccionado)
    replicador.ejecutar_replicacion()

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print(f"\n\n{Color.YELLOW}⚠{Color.END} Replicación cancelada.")
    except Exception as e:
        print(f"\n{Color.RED}✗{Color.END} Error inesperado: {e}")
        import traceback
        traceback.print_exc()